import re
import openpyxl

from io import BytesIO
from docx import Document
from .logger import get_logger
import zipfile
import xml.etree.ElementTree as ET

logger = get_logger(__name__)

def check_pii(text: str) -> bool:
    # 2.4 PII REGEX
    patterns = [
        (r"([A-Za-z0~9._%+-]+)@([A-Za-z0~9.-]+\.[A-Za-z]{2,})", "email"),
        (r"0\d{1,2}[- ]\d{3,4}[- ]\d{4}", "phone"),
        # (r"^(.*\d+(?:-\d+)?)(?:\s+(.*))?$", "address"),
        (r"([0-9]{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12][0-9]|3[01]))[-+ *. ]([5-8][0-9]{6})", "foreigner_id"),
        (r"\b\d{4}[- ]?(06|18)[- ]?\d{6}\b", "account_number_1"),
        (r"\b\d{4}(06|18)[- ]?\d{2}[- ]?\d{4}\b", "account_number_2"),
        (r"\b\d{6}[- ]?\d{2}[- ]?\d{6}\b", "account_number_3"),
        (r"\b\d{4}15[- ]?\d{2}[- ]?\d{5}\b", "account_number_4"),
        (r"[1-2][0-9]-[0-9]{2}-[0-9]{6}-[0-9]{2}", "driver_license"),
        (r"(010[-+ .*,]\d{4}[-+ .*,]\d{4}|8210[-+ .*,]\d{4}[-+ .*,]\d{4})", "phone_number_1"),
        (r"[0공][일1][0공][- ][\d일이삼사오육칠팔구]{3,4}[- ][\d일이삼사오육칠팔구]{4}", "phone_number_2"),
        (r"(^|\s)([DdMmRrSsGgTt])\d{8}(\s|$)", "passport_1"),
        (r"(^|\s)([DdMmRrSsGgTt])\d{3}[A-Z]\d{4}(\s|$)", "passport_2"),
        (r"(^|\s)[A-Z]{2}\d{7}(\s|$)", "passport_3"),
        (r"(?:9\d{3}|4\d{3}|5[1-5]\d{2}|6(?:011|5\d{2})|3[47]\d{2}|3(?:0[0-5]|[68]\d)|(?:2131|1800|35\d{3}))([- .]?\d{4}){3,4}", "credit_card"),
        (r"(?<![\dA-Za-z])\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])[\+\-\*\_][1-4]\d{6}(?![\dA-Za-z])", "resident_id")
    ]
    for pattern, name in patterns:
        match_result = re.search(pattern, text)
        if match_result:
            if name == "address" and match_result.group(2):
                return True
            if name != "address":
                return True
    return False



def mask_pii(text: str) -> bool:
    # 2.4 PII REGEX
    patterns = [
        (r"([A-Za-z0~9._%+-]+)@([A-Za-z0~9.-]+\.[A-Za-z]{2,})", "email"),
        (r"0\d{1,2}[- ]\d{3,4}[- ]\d{4}", "phone"),
        (r"([0-9]{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12][0-9]|3[01]))[-+ *. ]([5-8][0-9]{6})", "foreigner_id"),
        (r"\b\d{4}[- ]?(06|18)[- ]?\d{6}\b", "account_number_1"),
        (r"\b\d{4}(06|18)[- ]?\d{2}[- ]?\d{4}\b", "account_number_2"),
        (r"\b\d{6}[- ]?\d{2}[- ]?\d{6}\b", "account_number_3"),
        (r"\b\d{4}15[- ]?\d{2}[- ]?\d{5}\b", "account_number_4"),
        (r"[1-2][0-9]-[0-9]{2}-[0-9]{6}-[0-9]{2}", "driver_license"),
        (r"(010[-+ .*,]\d{4}[-+ .*,]\d{4}|8210[-+ .*,]\d{4}[-+ .*,]\d{4})", "phone_number_1"),
        (r"[0공][일1][0공][- ][\d일이삼사오육칠팔구]{3,4}[- ][\d일이삼사오육칠팔구]{4}", "phone_number_2"),
        (r"(^|\s)([DdMmRrSsGgTt])\d{8}(\s|$)", "passport_1"),
        (r"(^|\s)([DdMmRrSsGgTt])\d{3}[A-Z]\d{4}(\s|$)", "passport_2"),
        (r"(^|\s)[A-Z]{2}\d{7}(\s|$)", "passport_3"),
        (r"(?:9\d{3}|4\d{3}|5[1-5]\d{2}|6(?:011|5\d{2})|3[47]\d{2}|3(?:0[0-5]|[68]\d)|(?:2131|1800|35\d{3}))([- .]?\d{4}){3,4}", "credit_card"),
        (r"(?<![\dA-Za-z])\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])[\+\-\*\_][1-4]\d{6}(?![\dA-Za-z])", "resident_id")
    ]

    mask_char = "*"
    def masking(m):
        logger.info(f"마스킹 대상 발견: {m.group()}")
        return mask_char * len(m.group())
        
    for pat, name in patterns:
        text = re.sub(pat, masking, text)

    return text

def mask_md_bytes(data: bytes, encoding="utf-8") -> bytes:
    check_pii = False
    text = data.decode(encoding, errors="strict")
    masked = mask_pii(text)
    
    if masked != text:
        check_pii = True
    return masked.encode(encoding), check_pii

def mask_docx_bytes(data: bytes) -> bytes:
    inp = BytesIO(data)
    doc = Document(inp)
    check_pii = False
    
    for p in doc.paragraphs:
        for r in p.runs:
            mask_text = mask_pii(r.text)
            if r.text != mask_text:
                check_pii = True
            r.text = mask_text

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        mask_text = mask_pii(r.text)
                        if r.text != mask_text:
                            check_pii = True
                        r.text = mask_text
    
    out = BytesIO()
    doc.save(out)

    return out.getvalue(), check_pii

def mask_xlsx_bytes(data: bytes):
    # SpreadsheetML namespace
    NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    NS = {"main": NS_MAIN}

    inp = BytesIO(data)
    zin = zipfile.ZipFile(inp, "r")

    out_buf = BytesIO()
    zout = zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED)

    check_pii = False
    cache = {}

    def mask_cached(s: str) -> str:
        if not s:
            return s
        v = cache.get(s)
        if v is not None:
            return v
        m = mask_pii(s)
        cache[s] = m
        return m

    # 1) 수정 대상 파일들 이름 수집
    shared_path = "xl/sharedStrings.xml"
    worksheet_paths = [name for name in zin.namelist()
                       if name.startswith("xl/worksheets/") and name.endswith(".xml")]

    # 2) sharedStrings.xml 처리 (공유 문자열 테이블)
    if shared_path in zin.namelist():
        xml_bytes = zin.read(shared_path)
        root = ET.fromstring(xml_bytes)
        for t in root.findall(".//main:t", NS):
            if t.text:
                new_text = mask_cached(t.text)
                if new_text != t.text:
                    t.text = new_text
                    check_pii = True
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        zout.writestr(shared_path, new_xml)

    # 3) 각 워크시트의 inlineStr 처리 (sharedStrings를 쓰지 않는 셀들)
    for ws_path in worksheet_paths:
        xml_bytes = zin.read(ws_path)
        # 인라인 문자열이 없으면 파싱 비용 절약
        if b"<is" not in xml_bytes:
            # 그대로 복사
            zout.writestr(ws_path, xml_bytes)
            continue
        root = ET.fromstring(xml_bytes)
        modified = False
        # <is> 아래의 <t> 모두 수정 (리치 텍스트(<r>/<t>) 포함)
        for t in root.findall(".//main:is//main:t", NS):
            if t.text:
                new_text = mask_cached(t.text)
                if new_text != t.text:
                    t.text = new_text
                    modified = True
        if modified:
            check_pii = True
            new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(ws_path, new_xml)
        else:
            zout.writestr(ws_path, xml_bytes)

    # 4) 나머지 파일 그대로 복사
    to_skip = set([shared_path] + worksheet_paths)
    for name in zin.namelist():
        if name in to_skip:
            continue
        zout.writestr(name, zin.read(name))

    zout.close()
    return out_buf.getvalue(), check_pii
